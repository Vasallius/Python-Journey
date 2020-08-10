# Custom Invitation as Word Documents

import docx

# Open guests text file
guests = open('guests.txt')
guest_list = []

# Add each guest to guest_list
for line in guests.readlines():
    line = line.rstrip()
    guest_list.append(line)

# Load the document with our styles
doc = docx.Document('style.docx')
# Add the paragraph and set corresponding styles
for index, guest in enumerate(guest_list):
    doc.add_paragraph(
        'It would be a pleasure to have the company of').style = 'inv-heading'
    doc.add_paragraph(guest).style = 'name'
    doc.add_paragraph(
        'At 11010 Memory Lane on the Evening of').style = 'inv-heading'
    doc.add_paragraph('April 1st').style = 'inv-date'
    doc.add_paragraph('At 7 o’clock').style = 'inv-heading'
    # Don't add page break for the last person
    if index != len(guest_list)-1:
        doc.paragraphs[len(doc.paragraphs) -
                       1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    else:
        pass

doc.save('invitations.docx')
